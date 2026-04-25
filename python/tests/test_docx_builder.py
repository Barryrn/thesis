"""Tests for the HKA .docx thesis builder."""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from io import BytesIO
from docx import Document
from docx.shared import Cm, Pt

import docx_builder


def _minimal_data(**overrides) -> dict:
    """Return minimal thesis data for testing."""
    data = {
        "metadata": {
            "studentName": "Max Mustermann",
            "titleDE": "Testarbeit",
            "matrikelnummer": "12345",
            "studiengang": "Wirtschaftsinformatik",
            "erstpruefer": "Prof. Dr. Schmidt",
            "zweitpruefer": "Prof. Dr. Müller",
        },
        "sections": [],
        "sectionContents": {},
        "figures": [],
        "sources": [],
        "papers": {},
    }
    data.update(overrides)
    return data


class TestBuildEmpty:
    """Test building a document with minimal data."""

    def test_build_empty_document(self):
        """Building with no sections should produce a valid .docx."""
        data = _minimal_data()
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()
        assert doc is not None
        # Should be saveable without error
        buffer = BytesIO()
        doc.save(buffer)
        assert buffer.tell() > 0

    def test_title_page_has_metadata(self):
        """Title page should contain student name and thesis title."""
        data = _minimal_data()
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()

        # Collect all paragraph text
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Max Mustermann" in all_text
        assert "Testarbeit" in all_text
        assert "Masterarbeit" in all_text

    def test_declaration_text(self):
        """Eidesstattliche Erklärung should have exact HKA wording."""
        data = _minimal_data()
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "selbstständig und nur unter Benutzung" in all_text
        assert "keiner anderen Prüfungsbehörde vorgelegen" in all_text


class TestPageLayout:
    """Test page margins and layout."""

    def test_margins(self):
        """Page margins should match HKA spec."""
        data = _minimal_data()
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()

        section = doc.sections[0]
        # Margins should be approximately correct (±1mm tolerance)
        assert abs(section.top_margin - Cm(2.5)) < Cm(0.1)
        assert abs(section.bottom_margin - Cm(2.5)) < Cm(0.1)
        assert abs(section.left_margin - Cm(3)) < Cm(0.1)
        assert abs(section.right_margin - Cm(2.5)) < Cm(0.1)


class TestWithSections:
    """Test building with actual section content."""

    def test_sections_with_plain_text(self):
        """Sections should appear as headings with body text."""
        data = _minimal_data(
            sections=[
                {"_id": "s1", "title": "Einleitung", "orderNumber": "1", "depth": 0, "outlineVersion": 1},
                {"_id": "s2", "title": "Grundlagen", "orderNumber": "2", "depth": 0, "outlineVersion": 1},
            ],
            sectionContents={
                "s1": "This is the introduction.",
                "s2": "Background information here.",
            },
        )
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "1 Einleitung" in all_text
        assert "2 Grundlagen" in all_text
        assert "introduction" in all_text
        assert "Background" in all_text

    def test_sections_with_citations(self):
        """Citation markers should produce footnote-like text."""
        data = _minimal_data(
            sections=[
                {"_id": "s1", "title": "Test", "orderNumber": "1", "depth": 0, "outlineVersion": 1},
            ],
            sectionContents={
                "s1": "Text {{cite:p1::indirect::S. 3}} more.",
            },
            sources=[
                {"paperId": "p1", "kuerzel": "KR09", "sourceType": "book"},
            ],
            papers={
                "p1": {"_id": "p1", "title": "Test Paper", "authors": ["Author"], "year": 2009},
            },
        )
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()

        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Should contain the HKA footnote text
        assert "Vgl. KR09, S. 3." in all_text

    def test_sections_with_formulas(self):
        """Formula markers should be processed without errors."""
        data = _minimal_data(
            sections=[
                {"_id": "s1", "title": "Math", "orderNumber": "1", "depth": 0, "outlineVersion": 1},
            ],
            sectionContents={
                "s1": "The equation $x^2$ is simple.",
            },
        )
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()
        # Should not crash and produce valid document
        buffer = BytesIO()
        doc.save(buffer)
        assert buffer.tell() > 0


class TestBibliography:
    """Test bibliography generation."""

    def test_bibliography_sorted(self):
        """Bibliography entries should be sorted alphabetically."""
        data = _minimal_data(
            sources=[
                {"paperId": "p2", "kuerzel": "ZZ20", "sourceType": "book"},
                {"paperId": "p1", "kuerzel": "AA10", "sourceType": "book"},
            ],
            papers={
                "p1": {"_id": "p1", "title": "Alpha Paper", "authors": ["Alice Author"], "year": 2010},
                "p2": {"_id": "p2", "title": "Zeta Paper", "authors": ["Zoe Zebra"], "year": 2020},
            },
        )
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()

        all_text = "\n".join(p.text for p in doc.paragraphs)
        # AA10 should appear before ZZ20
        pos_aa = all_text.find("[AA10]")
        pos_zz = all_text.find("[ZZ20]")
        assert pos_aa < pos_zz


class TestFontStyles:
    """Test that font styles match HKA spec."""

    def test_normal_style(self):
        """Normal style should be Times New Roman 12pt."""
        data = _minimal_data()
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()

        style = doc.styles["Normal"]
        assert style.font.name == "Times New Roman"
        assert style.font.size == Pt(12)

    def test_heading1_style(self):
        """Heading 1 should be Arial 16pt bold."""
        data = _minimal_data()
        builder = docx_builder.HKADocxBuilder(data, {})
        doc = builder.build()

        style = doc.styles["Heading 1"]
        assert style.font.name == "Arial"
        assert style.font.size == Pt(16)
        assert style.font.bold is True
