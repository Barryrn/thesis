"""HKA-compliant .docx thesis builder.

Constructs a complete Word document from thesis data fetched from Convex,
following HKA Wirtschaftswissenschaften formatting requirements:
- Times New Roman 12pt, 1.5 spacing, justified
- Headings: sans-serif bold (16/14/12pt)
- Footnotes: 10pt single spacing
- Margins: top/bottom 2.5cm, left 3cm, right 2.5cm
- Pagination: no number on title page, roman for front matter, arabic for main
"""

from __future__ import annotations

import io
import logging
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree

import formula_converter

logger = logging.getLogger(__name__)

# ── Default layout settings (mirrors frontend DEFAULT_LAYOUT) ───────

DEFAULT_LAYOUT: dict = {
    "fontSize": 12,
    "lineSpacing": 1.5,
    "margins": {"top": 2.5, "bottom": 2.5, "left": 3, "right": 2.5},
    "pageNumbering": {"position": "header", "format": "automatic", "startPage": 1},
}

# ── Citation marker regex (mirrors frontend) ─────────────────────────

CITE_PRIMARY_RE = re.compile(
    r"\{\{cite:([^:]+)::(direct|indirect)::([^}]+)\}\}"
)
CITE_SECONDARY_RE = re.compile(
    r"\{\{cite:([^:]+)::(direct|indirect)::([^:]+)::via:([^:]+)::([^}]+)\}\}"
)
CITE_ANY_RE = re.compile(r"\{\{cite:[^}]+\}\}")

# Figure marker regex
FIG_MARKER_RE = re.compile(r"\{\{fig:([^}]+)\}\}")

# Formula regex (display then inline)
DISPLAY_MATH_RE = re.compile(r"\$\$([^$]+?)\$\$")
INLINE_MATH_RE = re.compile(r"(?<!\$)\$(?!\$)(\S[^$]*?\S|\S)\$(?!\$)")

# Combined marker regex for splitting body text
ALL_MARKERS_RE = re.compile(
    r"(\{\{cite:[^}]+\}\}|\{\{fig:[^}]+\}\}|\$\$[^$]+?\$\$|(?<!\$)\$(?!\$)(?:\S[^$]*?\S|\S)\$(?!\$))"
)

# ── HKA Eidesstattliche Erklärung ────────────────────────────────────

DECLARATION_TEXT = (
    "Hiermit erkläre ich, dass ich die vorliegende Arbeit selbstständig "
    "und nur unter Benutzung der angegebenen Quellen und Hilfsmittel "
    "angefertigt habe. Alle Textstellen, die wörtlich oder sinngemäß aus "
    "veröffentlichten oder nicht veröffentlichten Quellen entnommen wurden, "
    "sind als solche kenntlich gemacht. Die Arbeit hat in gleicher oder "
    "ähnlicher Form keiner anderen Prüfungsbehörde vorgelegen."
)


class HKADocxBuilder:
    """Builds an HKA-compliant .docx thesis document."""

    def __init__(self, data: dict, figure_bytes: dict[str, bytes]):
        self.data = data
        self.figure_bytes = figure_bytes
        self.doc = Document()
        self.footnote_counter = 0
        self._footnotes_part = None
        self._footnotes_xml = None

        # Resolve layout settings — deep-merge user overrides over defaults.
        raw_layout = (data.get("metadata") or {}).get("layoutSettings") or {}
        self.layout: dict = {
            **DEFAULT_LAYOUT,
            **{k: v for k, v in raw_layout.items() if k in DEFAULT_LAYOUT and not isinstance(v, dict)},
            "margins": {**DEFAULT_LAYOUT["margins"], **(raw_layout.get("margins") or {})},
            "pageNumbering": {**DEFAULT_LAYOUT["pageNumbering"], **(raw_layout.get("pageNumbering") or {})},
        }

        # Build lookup maps
        self.source_map: dict[str, dict] = {}
        for src in data.get("sources", []):
            self.source_map[src["paperId"]] = src

        self.paper_map: dict[str, dict] = data.get("papers", {})

        # Build figure map
        self.figure_map: dict[str, dict] = {}
        for fig in data.get("figures", []):
            self.figure_map[fig["_id"]] = fig

        # Compute figure numbers
        self.figure_numbers: dict[str, int] = {}
        self._compute_figure_numbers()

    def build(self) -> Document:
        """Build the complete thesis document and return it."""
        self._setup_styles()
        self._setup_page_layout()
        self._add_title_page()
        self._add_abstracts()
        self._add_toc()
        self._add_list_of_figures()
        self._add_main_content()
        self._add_bibliography()
        self._add_declaration()
        return self.doc

    # ── Styles ────────────────────────────────────────────────────────

    def _setup_styles(self):
        """Configure document styles using resolved layout settings."""
        font_size = self.layout["fontSize"]
        line_spacing = self.layout["lineSpacing"]

        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(font_size)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Heading sizes scale relative to body font: +4, +2, +0
        for level, offset in [(1, 4), (2, 2), (3, 0)]:
            h_style = self.doc.styles[f"Heading {level}"]
            h_style.font.name = "Arial"
            h_style.font.size = Pt(font_size + offset)
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0, 0, 0)

        # Footnote text style: 10pt Times New Roman, single spacing
        try:
            fn_text_style = self.doc.styles.add_style('FootnoteText', WD_STYLE_TYPE.PARAGRAPH)
            fn_text_style.font.name = 'Times New Roman'
            fn_text_style.font.size = Pt(10)
            fn_text_style.paragraph_format.line_spacing = 1.0
        except ValueError:
            pass  # Style already exists

        # Footnote reference style: superscript
        try:
            fn_ref_style = self.doc.styles.add_style('FootnoteReference', WD_STYLE_TYPE.CHARACTER)
            fn_ref_style.font.superscript = True
            fn_ref_style.font.size = Pt(10)
        except ValueError:
            pass  # Style already exists

    def _setup_page_layout(self):
        """Configure page margins on the initial (title page) section.

        The title page section intentionally has no header so that no
        page number is visible on the Deckblatt.
        """
        self._apply_margins(self.doc.sections[0])

    def _apply_margins(self, section):
        """Apply configured margins to a document section."""
        m = self.layout["margins"]
        section.top_margin = Cm(m["top"])
        section.bottom_margin = Cm(m["bottom"])
        section.left_margin = Cm(m["left"])
        section.right_margin = Cm(m["right"])

    def _add_page_number_field(self, section):
        """Add a right-aligned PAGE field to the section header or footer.

        The position (header vs footer) is determined by layout settings.
        """
        position = self.layout["pageNumbering"]["position"]

        if position == "footer":
            container = section.footer
        else:
            container = section.header

        container.is_linked_to_previous = False
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = p.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        run._element.append(instr_text)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._element.append(fld_char_end)

    # ── Title page ────────────────────────────────────────────────────

    def _add_title_page(self):
        """Add the Deckblatt with HKA-required fields."""
        metadata = self.data.get("metadata") or {}

        # Thesis type
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(72)
        run = p.add_run("Masterarbeit")
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.bold = True

        # German title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(24)
        run = p.add_run(metadata.get("titleDE", ""))
        run.font.name = "Arial"
        run.font.size = Pt(18)
        run.bold = True

        # German subtitle
        if metadata.get("subtitleDE"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(metadata["subtitleDE"])
            run.font.size = Pt(14)
            run.italic = True

        # English title
        if metadata.get("titleEN"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(12)
            run = p.add_run(metadata["titleEN"])
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.bold = True

        # English subtitle
        if metadata.get("subtitleEN"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(metadata["subtitleEN"])
            run.font.size = Pt(12)
            run.italic = True

        # Vorlage text
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(24)
        run = p.add_run(
            "Masterarbeit vorgelegt zur Erlangung des Mastergrades "
            "der Hochschule Karlsruhe – Technik und Wirtschaft"
        )
        run.font.size = Pt(10)
        run.italic = True

        # Personal info table
        self.doc.add_paragraph()  # spacing
        info_fields = [
            ("Name", metadata.get("studentName", "")),
            ("Adresse", metadata.get("studentAddress", "")),
            ("E-Mail (privat)", metadata.get("studentEmail", "")),
            ("E-Mail (Hochschule)", metadata.get("universityEmail", "")),
            ("Matrikelnummer", metadata.get("matrikelnummer", "")),
            ("Studiengang", metadata.get("studiengang", "")),
            ("Erstprüfer", metadata.get("erstpruefer", "")),
            ("Zweitprüfer", metadata.get("zweitpruefer", "")),
            ("Abgabedatum", metadata.get("abgabedatum", "")),
        ]
        for label, value in info_fields:
            if value:
                p = self.doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.15
                run_label = p.add_run(f"{label}: ")
                run_label.bold = True
                run_label.font.size = Pt(10)
                run_value = p.add_run(value)
                run_value.font.size = Pt(10)

        # End title page section — the next section will be front matter
        # with roman numeral page numbering.
        front_matter = self.doc.add_section(WD_SECTION.NEW_PAGE)
        self._apply_margins(front_matter)
        self._add_page_number_field(front_matter)

        # Page numbering format for front matter — respects layout settings.
        pn = self.layout["pageNumbering"]
        front_fmt = {
            "automatic": "lowerRoman",
            "allArabic": "decimal",
            "allRoman": "lowerRoman",
        }.get(pn["format"], "lowerRoman")
        pgNumFmt = OxmlElement("w:pgNumType")
        pgNumFmt.set(qn("w:fmt"), front_fmt)
        pgNumFmt.set(qn("w:start"), str(pn["startPage"]))
        front_matter._sectPr.append(pgNumFmt)

    # ── Abstracts ─────────────────────────────────────────────────────

    def _add_abstracts(self):
        """Add Kurzfassung (DE) and Abstract (EN) with keywords.

        Only adds a trailing page break if at least one abstract was written,
        preventing a blank page when both abstracts are empty.
        """
        metadata = self.data.get("metadata") or {}
        has_abstract = False

        if metadata.get("abstractDE"):
            has_abstract = True
            self.doc.add_heading("Kurzfassung", level=1)
            self.doc.add_paragraph(metadata["abstractDE"])
            if metadata.get("keywordsDE"):
                p = self.doc.add_paragraph()
                run = p.add_run("Schlüsselbegriffe: ")
                run.bold = True
                run.font.size = Pt(10)
                run2 = p.add_run(", ".join(metadata["keywordsDE"]))
                run2.font.size = Pt(10)
                run2.italic = True

        if metadata.get("abstractEN"):
            has_abstract = True
            self.doc.add_heading("Abstract", level=1)
            self.doc.add_paragraph(metadata["abstractEN"])
            if metadata.get("keywordsEN"):
                p = self.doc.add_paragraph()
                run = p.add_run("Keywords: ")
                run.bold = True
                run.font.size = Pt(10)
                run2 = p.add_run(", ".join(metadata["keywordsEN"]))
                run2.font.size = Pt(10)
                run2.italic = True

        if has_abstract:
            self.doc.add_page_break()

    # ── TOC ───────────────────────────────────────────────────────────

    def _add_toc(self):
        """Add a Table of Contents that auto-updates in Word."""
        self.doc.add_heading("Inhaltsverzeichnis", level=1)

        # Insert TOC field code
        p = self.doc.add_paragraph()
        run = p.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._element.append(instr)

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._element.append(fld_sep)

        # Placeholder text
        run2 = p.add_run("(Update this field to see table of contents)")
        run2.font.color.rgb = RGBColor(128, 128, 128)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run2._element.append(fld_end)

        self.doc.add_page_break()

    # ── List of figures ───────────────────────────────────────────────

    def _add_list_of_figures(self):
        """Add Abbildungsverzeichnis if there are any figures."""
        if not self.figure_numbers:
            return

        self.doc.add_heading("Abbildungsverzeichnis", level=1)

        sections = self._sorted_sections()
        for section in sections:
            sec_figs = [
                f for f in self.data.get("figures", [])
                if f["sectionId"] == section["_id"]
            ]
            sec_figs.sort(key=lambda f: f["orderIndex"])
            for fig in sec_figs:
                num = self.figure_numbers.get(fig["_id"])
                if num is not None:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.15
                    run = p.add_run(f"Abbildung {num}: {fig['caption']}")
                    run.font.size = Pt(10)

        # No trailing page break — the main content section break handles the transition.

    # ── Main content ──────────────────────────────────────────────────

    def _add_main_content(self):
        """Add all thesis sections with rendered body content.

        Inserts a section break before the first chapter to switch from
        roman (front matter) to arabic (main content) page numbering.
        Top-level chapters (depth 0) each start on a new page.
        """
        sections = self._sorted_sections()
        contents = self.data.get("sectionContents", {})
        # Per-chapter equation counters for numbering like (2.1), (2.2), (3.1)
        equation_counters: dict[str, int] = {}

        # Start a new Word section for arabic page numbering (1, 2, 3...)
        main_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        self._apply_margins(main_section)
        self._add_page_number_field(main_section)

        # Page numbering format for main content — respects layout settings.
        pn = self.layout["pageNumbering"]
        main_fmt = {
            "automatic": "decimal",
            "allArabic": "decimal",
            "allRoman": "lowerRoman",
        }.get(pn["format"], "decimal")
        pgNumFmt = OxmlElement("w:pgNumType")
        pgNumFmt.set(qn("w:fmt"), main_fmt)
        pgNumFmt.set(qn("w:start"), str(pn["startPage"]))
        main_section._sectPr.append(pgNumFmt)

        # Determine the top-level depth dynamically (may be 0 or 1 depending
        # on how the outline was built in the editor).
        min_depth = min((s.get("depth", 0) for s in sections), default=0)

        first_chapter = True
        for section in sections:
            depth = section.get("depth", 0)

            # Top-level chapters start on a new page (except the first,
            # which already follows the section break above).
            if depth == min_depth and not first_chapter:
                self.doc.add_page_break()
            if depth == min_depth:
                first_chapter = False

            level = min(depth - min_depth + 1, 3)
            title = f"{section['orderNumber']} {section['title']}"
            self.doc.add_heading(title, level=level)

            body = contents.get(section["_id"], "")
            if body.strip():
                chapter_num = section["orderNumber"].split(".")[0]
                self._render_body(body, chapter_num, equation_counters)

    def _render_body(self, body: str, chapter_number: str = "", equation_counters: dict[str, int] | None = None):
        """Render body text, processing citation/figure/formula markers."""
        if equation_counters is None:
            equation_counters = {}
        # Split by all marker types while keeping the markers.
        # NOTE: footnote_counter is NOT reset here — it accumulates across
        # sections so numbering is continuous throughout the entire document.
        parts = ALL_MARKERS_RE.split(body)

        current_paragraph = None

        for part in parts:
            if not part:
                continue

            cite_match = CITE_ANY_RE.fullmatch(part)
            fig_match = FIG_MARKER_RE.fullmatch(part)
            display_math_match = DISPLAY_MATH_RE.fullmatch(part)
            inline_math_match = INLINE_MATH_RE.fullmatch(part)

            if cite_match:
                # Citation marker → Word footnote
                if current_paragraph is None:
                    current_paragraph = self.doc.add_paragraph()
                self._add_footnote_for_citation(current_paragraph, part)

            elif fig_match:
                # Figure marker → embedded image
                fig_id = fig_match.group(1)
                current_paragraph = None  # figures are block-level
                self._add_figure(fig_id)

            elif display_math_match:
                # Display formula → centered OMML with right-aligned equation number
                latex = display_math_match.group(1)
                current_paragraph = None
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                formula_converter.insert_formula(p, latex, display_mode=True)

                # Add per-chapter equation number, e.g. (2.1)
                eq_num = equation_counters.get(chapter_number, 1)
                equation_counters[chapter_number] = eq_num + 1
                eq_label = f"({chapter_number}.{eq_num})" if chapter_number else f"({eq_num})"

                # Right-aligned tab stop at right margin (~16cm) for the number
                pPr = p._element.get_or_add_pPr()
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9072')  # twips: ~16cm from left
                tabs.append(tab)
                pPr.append(tabs)

                run = p.add_run(f'\t{eq_label}')
                run.font.size = Pt(12)

            elif inline_math_match:
                # Inline formula → OMML within current paragraph
                latex = inline_math_match.group(1)
                if current_paragraph is None:
                    current_paragraph = self.doc.add_paragraph()
                formula_converter.insert_formula(current_paragraph, latex, display_mode=False)

            else:
                # Plain text — split by newlines for paragraphs
                lines = part.split("\n")
                for i, line in enumerate(lines):
                    if i > 0:
                        current_paragraph = None  # new paragraph on newline
                    if line.strip():
                        if current_paragraph is None:
                            current_paragraph = self.doc.add_paragraph()
                        current_paragraph.add_run(line)

    def _add_footnote_for_citation(self, paragraph, marker: str):
        """Add a Word footnote for a citation marker."""
        # Parse the marker
        sec_match = CITE_SECONDARY_RE.match(marker)
        pri_match = CITE_PRIMARY_RE.match(marker)

        if sec_match:
            paper_id = sec_match.group(1)
            cite_type = sec_match.group(2)
            page_ref = sec_match.group(3)
            sec_paper_id = sec_match.group(4)
            sec_page_ref = sec_match.group(5)
        elif pri_match:
            paper_id = pri_match.group(1)
            cite_type = pri_match.group(2)
            page_ref = pri_match.group(3)
            sec_paper_id = None
            sec_page_ref = None
        else:
            return

        # Build footnote text
        source = self.source_map.get(paper_id, {})
        kuerzel = source.get("kuerzel", "??")
        prefix = "Vgl. " if cite_type == "indirect" else ""

        if sec_paper_id and sec_page_ref:
            sec_source = self.source_map.get(sec_paper_id, {})
            sec_kuerzel = sec_source.get("kuerzel", "??")
            fn_text = f"{prefix}{kuerzel}, {page_ref} zitiert nach {sec_kuerzel}, {sec_page_ref}."
        else:
            fn_text = f"{prefix}{kuerzel}, {page_ref}."

        # Create footnote via XML manipulation
        self.footnote_counter += 1
        self._insert_footnote(paragraph, fn_text)

    def _ensure_footnotes_part(self):
        """Create or return the footnotes XML part for the document.

        Word requires a footnotes.xml part with separator footnotes (ids 0/1)
        before any real footnotes can be added. This method creates that part
        on first call and returns the cached XML root on subsequent calls.
        """
        if self._footnotes_xml is not None:
            return self._footnotes_xml

        # XML namespaces
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        }

        # Build footnotes root with required separator footnotes (ids 0 and 1)
        footnotes_root = etree.Element(qn('w:footnotes'), nsmap=nsmap)

        # Separator footnote (id=0) — the thin line above footnotes
        fn0 = etree.SubElement(footnotes_root, qn('w:footnote'))
        fn0.set(qn('w:type'), 'separator')
        fn0.set(qn('w:id'), '0')
        p0 = etree.SubElement(fn0, qn('w:p'))
        r0 = etree.SubElement(p0, qn('w:r'))
        sep = etree.SubElement(r0, qn('w:separator'))

        # Continuation separator (id=1)
        fn1 = etree.SubElement(footnotes_root, qn('w:footnote'))
        fn1.set(qn('w:type'), 'continuationSeparator')
        fn1.set(qn('w:id'), '1')
        p1 = etree.SubElement(fn1, qn('w:p'))
        r1 = etree.SubElement(p1, qn('w:r'))
        cont = etree.SubElement(r1, qn('w:continuationSeparator'))

        # Serialize to bytes
        blob = etree.tostring(footnotes_root, xml_declaration=True, encoding='UTF-8', standalone=True)

        # Create the Part
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
        part_name = PackURI('/word/footnotes.xml')
        footnotes_part = Part(part_name, content_type, blob, self.doc.part.package)

        # Register relationship
        self.doc.part.relate_to(footnotes_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')

        # Cache
        self._footnotes_part = footnotes_part
        self._footnotes_xml = footnotes_root
        return self._footnotes_xml

    def _insert_footnote(self, paragraph, text: str):
        """Insert a proper Word footnote with reference in body and text at page bottom."""
        footnotes_xml = self._ensure_footnotes_part()

        # Footnote IDs 0 and 1 are reserved for separators, so real footnotes start at 2
        fn_id = str(self.footnote_counter + 1)

        # 1. Add footnote body to footnotes.xml
        fn_elem = etree.SubElement(footnotes_xml, qn('w:footnote'))
        fn_elem.set(qn('w:id'), fn_id)

        p_elem = etree.SubElement(fn_elem, qn('w:p'))

        # Paragraph properties: FootnoteText style, single spacing
        pPr = etree.SubElement(p_elem, qn('w:pPr'))
        pStyle = etree.SubElement(pPr, qn('w:pStyle'))
        pStyle.set(qn('w:val'), 'FootnoteText')

        # Auto-number reference inside the footnote (the superscript number at the start)
        r_ref = etree.SubElement(p_elem, qn('w:r'))
        rPr_ref = etree.SubElement(r_ref, qn('w:rPr'))
        rStyle_ref = etree.SubElement(rPr_ref, qn('w:rStyle'))
        rStyle_ref.set(qn('w:val'), 'FootnoteReference')
        etree.SubElement(r_ref, qn('w:footnoteRef'))

        # Space + footnote text
        r_text = etree.SubElement(p_elem, qn('w:r'))
        rPr_text = etree.SubElement(r_text, qn('w:rPr'))
        sz = etree.SubElement(rPr_text, qn('w:sz'))
        sz.set(qn('w:val'), '20')  # 10pt
        szCs = etree.SubElement(rPr_text, qn('w:szCs'))
        szCs.set(qn('w:val'), '20')  # 10pt
        t = etree.SubElement(r_text, qn('w:t'))
        t.set(qn('xml:space'), 'preserve')
        t.text = f' {text}'

        # Update the underlying blob in the part (Part.blob is read-only,
        # so we write to the private _blob attribute directly).
        self._footnotes_part._blob = etree.tostring(
            footnotes_xml, xml_declaration=True, encoding='UTF-8', standalone=True
        )

        # 2. Add footnote reference in the body paragraph
        fn_ref_run = OxmlElement('w:r')
        fn_ref_rPr = OxmlElement('w:rPr')
        fn_ref_rStyle = OxmlElement('w:rStyle')
        fn_ref_rStyle.set(qn('w:val'), 'FootnoteReference')
        fn_ref_rPr.append(fn_ref_rStyle)
        fn_ref_run.append(fn_ref_rPr)

        fn_reference = OxmlElement('w:footnoteReference')
        fn_reference.set(qn('w:id'), fn_id)
        fn_ref_run.append(fn_reference)

        paragraph._element.append(fn_ref_run)

    def _add_figure(self, figure_id: str):
        """Add a figure with image and caption."""
        figure = self.figure_map.get(figure_id)
        num = self.figure_numbers.get(figure_id)

        if not figure or num is None:
            p = self.doc.add_paragraph()
            run = p.add_run("[Missing figure]")
            run.italic = True
            return

        img_bytes = self.figure_bytes.get(figure_id)
        if img_bytes:
            # Add image
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(io.BytesIO(img_bytes), width=Cm(14))
            except Exception as e:
                logger.warning(f"Failed to embed figure {figure_id}: {e}")
                run.add_text("[Image could not be embedded]")
        else:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("[Image not available]")
            run.italic = True

        # Caption
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(f"Abbildung {num}: {figure['caption']}")
        run.font.size = Pt(10)
        run.italic = True

    # ── Bibliography ──────────────────────────────────────────────────

    def _add_bibliography(self):
        """Add Literaturverzeichnis with formatted entries."""
        sources = self.data.get("sources", [])
        if not sources:
            return

        self.doc.add_page_break()
        self.doc.add_heading("Literaturverzeichnis", level=1)

        entries = []
        for src in sources:
            paper = self.paper_map.get(src["paperId"])
            if not paper:
                continue

            kuerzel = src.get("kuerzel", "??")
            author_str = self._format_authors(paper.get("authors", []))
            year = paper.get("year")
            year_str = f"({year})" if year else "(o.J.)"
            title = paper.get("title", "")

            text = f"{author_str} {year_str}. {title}."

            source_type = src.get("sourceType", "book")
            if source_type == "book":
                loc = src.get("publisherLocation", "")
                pub = src.get("publisher", "")
                if loc and pub:
                    text += f" {loc}: {pub}."
            elif source_type == "journalArticle":
                if src.get("journalName"):
                    text += f" {src['journalName']}."
                if src.get("volume"):
                    text += f" Jahrgang {src['volume']}"
                if src.get("issue"):
                    text += f", Heft {src['issue']}"
                text += "."
            elif source_type == "internetSource":
                if src.get("accessDate"):
                    text += f" Abgerufen am {src['accessDate']}"
                if src.get("url"):
                    text += f" von {src['url']}"
                text += "."

            entries.append((kuerzel, text))

        # Sort alphabetically by formatted text
        entries.sort(key=lambda e: e[1].lower())

        for kuerzel, text in entries:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(-1.5)
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(f"[{kuerzel}] ")
            run.bold = True
            p.add_run(text)

    # ── Declaration ───────────────────────────────────────────────────

    def _add_declaration(self):
        """Add Eidesstattliche Erklärung."""
        self.doc.add_page_break()
        self.doc.add_heading("Eidesstattliche Erklärung", level=1)
        self.doc.add_paragraph(DECLARATION_TEXT)

        # Signature line
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run("_" * 30 + "          " + "_" * 30)
        p2 = self.doc.add_paragraph()
        run1 = p2.add_run("Ort, Datum")
        run1.font.size = Pt(10)
        p2.add_run("                              ")
        run2 = p2.add_run("Unterschrift")
        run2.font.size = Pt(10)

    # ── Helpers ───────────────────────────────────────────────────────

    def _sorted_sections(self) -> list[dict]:
        """Return outline sections sorted by orderNumber (numeric-aware).

        Uses tuple-of-ints sorting so "2" < "10" and "1.2" < "1.10",
        matching the frontend's localeCompare with { numeric: true }.
        """
        sections = self.data.get("sections", [])
        return sorted(
            sections,
            key=lambda s: [int(x) for x in s.get("orderNumber", "0").split(".")],
        )

    def _compute_figure_numbers(self):
        """Assign sequential figure numbers across the document."""
        sections = self._sorted_sections()
        counter = 1
        for section in sections:
            sec_figs = [
                f for f in self.data.get("figures", [])
                if f["sectionId"] == section["_id"]
            ]
            sec_figs.sort(key=lambda f: f["orderIndex"])
            for fig in sec_figs:
                self.figure_numbers[fig["_id"]] = counter
                counter += 1

    def _format_authors(self, authors: list[str]) -> str:
        """Format author list for bibliography entries."""
        if not authors:
            return "Unbekannt"
        formatted = []
        for a in authors:
            parts = a.strip().split()
            if len(parts) == 1:
                formatted.append(parts[0])
            else:
                surname = parts[-1]
                initials = " ".join(p[0] + "." for p in parts[:-1])
                formatted.append(f"{surname}, {initials}")

        if len(formatted) == 1:
            return formatted[0]
        if len(formatted) == 2:
            return f"{formatted[0]} und {formatted[1]}"
        return ", ".join(formatted[:-1]) + " und " + formatted[-1]
